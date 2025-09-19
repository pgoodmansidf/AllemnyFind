# app/services/smartmatch_service.py
import logging
import asyncio
import re
import json
import hashlib
from typing import Dict, Any, List, Optional, AsyncGenerator
from datetime import datetime
from uuid import uuid4
import numpy as np
from io import BytesIO

from sqlalchemy import text, and_, or_, func
from sqlalchemy.orm import Session
from sqlalchemy.exc import SQLAlchemyError
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain_ollama import OllamaEmbeddings

# Fixed import - using the correct class name
from app.services.langchain_processor import AgenticDocumentProcessor
from app.models.document import Document, DocumentChunk
from app.core.config import settings

logger = logging.getLogger(__name__)

class SmartMatchService:
    """Service for intelligent document matching using AI"""
    
    def __init__(self, groq_api_key: str, db: Session):
        self.db = db
        self.llm = ChatGroq(
            api_key=groq_api_key,
            model=settings.groq_model,
            temperature=0.3,
            streaming=True
        )
        
        self.embeddings = OllamaEmbeddings(
            base_url=settings.ollama_base_url,
            model="nomic-embed-text"
        )
        
        # Fixed: Using AgenticDocumentProcessor instead of DocumentProcessor
        self.processor = AgenticDocumentProcessor()
        
        # Analysis tasks storage (in production, use Redis)
        self.analysis_tasks = {}

        # Add this configuration
        self.max_results = 10  # Maximum number of similar documents to return
        
        # Extraction prompt
        self.extraction_prompt = PromptTemplate(
            template="""Analyze the following document content and extract key information.

Document Content:
{content}

Please extract and provide in JSON format:
1. "products": List of main products or services mentioned (max 10)
2. "sau_numbers": List of SAU numbers (format: SAU followed by numbers, e.g., SAU1234, SAU 5678, SAU-9012, SAU #3456)
3. "main_city": The primary city or region mentioned
4. "main_company": The main company or organization discussed
5. "industry": The industry this document relates to
6. "sector": The specific sector within the industry
7. "summary": A brief 2-3 sentence summary of the document

Be precise and extract only what is explicitly mentioned in the document.
Return ONLY valid JSON format.""",
            input_variables=["content"]
        )
        
    async def start_analysis(self, file, user_id: int, file_content: bytes) -> str:
        """Start document analysis task"""
        task_id = str(uuid4())
        
        # Store task info
        self.analysis_tasks[task_id] = {
            "user_id": user_id,
            "status": "processing",
            "started_at": datetime.now(),
            "file_name": file.filename,
            "file_size": len(file_content),
            "file_content": file_content
        }
        
        # Start processing in background
        asyncio.create_task(self._process_document(task_id, file.filename, file_content))
        
        return task_id
    
    async def _process_document(self, task_id: str, filename: str, content: bytes):
        """Process uploaded document"""
        try:
            # Extract text based on file type
            text = ""
            
            if filename.lower().endswith('.pdf'):
                # Process PDF using the AgenticDocumentProcessor
                try:
                    # Create a ProcessedDocument object
                    from app.services.langchain_processor import ProcessedDocument, DocumentType
                    from pathlib import Path
                    
                    # Create a temporary processed document
                    temp_doc = ProcessedDocument(
                        file_path=Path(filename),
                        doc_type=DocumentType.PDF,
                        content="",
                        metadata={},
                        chunks=[],
                        tables=[]
                    )
                    
                    # Extract text using processor methods
                    import pdfplumber
                    from io import BytesIO
                    
                    with pdfplumber.open(BytesIO(content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except Exception as e:
                    logger.error(f"Error extracting PDF text: {e}")
                    # Fallback to basic extraction
                    text = content.decode('utf-8', errors='ignore')
                    
            elif filename.lower().endswith(('.docx', '.doc')):
                # Process Word documents
                try:
                    from docx import Document as DocxDocument
                    from io import BytesIO
                    
                    doc = DocxDocument(BytesIO(content))
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                except Exception as e:
                    logger.error(f"Error extracting Word text: {e}")
                    text = content.decode('utf-8', errors='ignore')
                    
            elif filename.lower().endswith(('.xlsx', '.xls')):
                # Process Excel
                try:
                    import pandas as pd
                    from io import BytesIO
                    
                    excel_file = BytesIO(content)
                    # Read all sheets
                    excel_data = pd.read_excel(excel_file, sheet_name=None)
                    
                    for sheet_name, df in excel_data.items():
                        text += f"\nSheet: {sheet_name}\n"
                        text += df.to_string() + "\n"
                except Exception as e:
                    logger.error(f"Error extracting Excel text: {e}")
                    text = content.decode('utf-8', errors='ignore')
            else:
                # Plain text
                text = content.decode('utf-8', errors='ignore')
            
            # Clean the text
            text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
            text = text.strip()
            
            # Limit text for analysis (first 10000 chars)
            text = text[:10000] if len(text) > 10000 else text
            
            # Store extracted text
            self.analysis_tasks[task_id]["extracted_text"] = text
            self.analysis_tasks[task_id]["status"] = "analyzing"
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            self.analysis_tasks[task_id]["status"] = "error"
            self.analysis_tasks[task_id]["error"] = str(e)
    
    async def analyze_document_stream(
        self, 
        task_id: str, 
        user_id: int
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """Stream analysis progress and results"""
        
        # Wait for text extraction
        max_wait = 30  # 30 seconds timeout
        waited = 0
        
        while waited < max_wait:
            if task_id not in self.analysis_tasks:
                yield {
                    "type": "error",
                    "message": "Task not found"
                }
                return
            
            task = self.analysis_tasks[task_id]
            
            if task["user_id"] != user_id:
                yield {
                    "type": "error",
                    "message": "Unauthorized"
                }
                return
            
            if "extracted_text" in task:
                break
            
            if task["status"] == "error":
                yield {
                    "type": "error",
                    "message": task.get("error", "Processing failed")
                }
                return
            
            await asyncio.sleep(1)
            waited += 1
        
        if waited >= max_wait:
            yield {
                "type": "error",
                "message": "Processing timeout"
            }
            return
        
        try:
            # Stage 1: Extract metadata
            yield {
                "type": "stage",
                "message": "Extracting document metadata..."
            }
            
            text = task["extracted_text"]
            
            # Use Groq to extract metadata
            extraction_prompt = self.extraction_prompt.format(content=text)
            
            metadata_response = ""
            async for chunk in self.llm.astream(extraction_prompt):
                if chunk.content:
                    metadata_response += chunk.content
            
            # Parse JSON response
            try:
                # Clean response to get JSON
                json_match = re.search(r'\{.*\}', metadata_response, re.DOTALL)
                if json_match:
                    metadata = json.loads(json_match.group())
                else:
                    raise ValueError("No JSON found in response")
            except:
                # Fallback metadata
                metadata = {
                    "products": [],
                    "sau_numbers": [],
                    "main_city": "",
                    "main_company": "",
                    "industry": "General",
                    "sector": "Unspecified",
                    "summary": "Document analysis completed"
                }
            
            # Clean and extract SAU numbers
            metadata["sau_numbers"] = self._extract_sau_numbers(text)
            
            yield {
                "type": "metadata",
                "data": metadata
            }
            
            # Stage 2: Generate embedding
            yield {
                "type": "stage",
                "message": "Generating document embedding..."
            }
            
            embedding = await self.embeddings.aembed_query(text[:2000])
            
            # Stage 3: Find similar documents
            yield {
                "type": "stage",
                "message": "Searching for similar documents..."
            }
            
            similar_docs = await self._find_similar_documents(
                embedding=embedding,
                metadata=metadata
            )
            
            # Stage 4: Calculate similarity scores
            yield {
                "type": "stage",
                "message": "Calculating similarity scores..."
            }
            
            # Calculate detailed similarity scores
            scored_docs = self._calculate_similarity_scores(
                similar_docs,
                metadata,
                embedding
            )
            
            # Prepare final result
            from datetime import timezone

            # Make sure started_at is timezone-aware
            started_at = task["started_at"]
            if isinstance(started_at, datetime):
                if started_at.tzinfo is None:
                    started_at = started_at.replace(tzinfo=timezone.utc)
                current_time = datetime.now(timezone.utc)
                analysis_time = (current_time - started_at).total_seconds()
            else:
                analysis_time = 0

            # After calculating similarity scores, filter by minimum threshold
            MIN_SIMILARITY_THRESHOLD = 0.5  # Only show matches above 30% similarity

            # Filter out low-quality matches
            scored_docs = [doc for doc in scored_docs if doc['similarity_score'] >= MIN_SIMILARITY_THRESHOLD]
    

            result = {
                "uploaded_document": metadata,
                "similar_documents": scored_docs[:self.max_results],  # Use configurable limit
                "total_matches": len(scored_docs),
                "analysis_time": analysis_time  # This is now a float, not a timedelta
            }
            
            # Store result
            self.analysis_tasks[task_id]["result"] = result
            self.analysis_tasks[task_id]["status"] = "completed"
            
            yield {
                "type": "result",
                "result": result
            }
            
        except Exception as e:
            logger.error(f"Error in analysis stream: {e}")
            yield {
                "type": "error",
                "message": str(e)
            }
    
    def _extract_sau_numbers(self, text: str) -> List[str]:
        """Extract SAU numbers from text"""
        # Pattern for SAU numbers
        patterns = [
            r'SAU\s*#?\s*(\d+)',
            r'SAU-(\d+)',
            r'SAU(\d+)',
        ]
        
        sau_numbers = set()
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                sau_numbers.add(f"SAU{match}")
        
        return list(sau_numbers)
    
    async def _find_similar_documents(
        self, 
        embedding: List[float],
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Find similar documents in database"""
        try:
            # Rollback any existing transaction
            self.db.rollback()
            
            embedding_str = '[' + ','.join(map(str, embedding)) + ']'
            
            # Updated query to match your database schema
            sql = text("""
                WITH similar_docs AS (
                    SELECT DISTINCT
                        d.id,
                        d.filename,
                        d.title,
                        d.summary,
                        d.product_tags,
                        d.main_tag,
                        d.file_type,
                        d.file_size,
                        d.modification_date,
                        d.project_number,
                        d.main_city,
                        d.companies AS main_companies,
                        MIN(dc.embedding <=> CAST(:embedding AS vector)) as vector_distance
                    FROM documents d
                    LEFT JOIN document_chunks dc ON d.id = dc.document_id
                    WHERE dc.embedding IS NOT NULL
                    GROUP BY d.id, d.filename, d.title, d.summary, d.product_tags,
                             d.main_tag, d.file_type, d.file_size, d.modification_date,
                             d.project_number, d.main_city, d.companies
                    ORDER BY vector_distance
                    LIMIT 100
                )
                SELECT 
                    sd.*,
                    COUNT(DISTINCT ds.user_id) as star_count,
                    COUNT(DISTINCT contrib.id) as contribution_count
                FROM similar_docs sd
                LEFT JOIN document_stars ds ON sd.id = ds.document_id
                LEFT JOIN document_contributions contrib ON sd.id = contrib.document_id
                GROUP BY 
                    sd.id, sd.filename, sd.title, sd.summary, 
                    sd.product_tags, sd.main_tag, sd.file_type, sd.file_size,
                    sd.modification_date, sd.vector_distance, sd.project_number,
                    sd.main_city, sd.main_companies
                ORDER BY sd.vector_distance
            """)
            
            result = self.db.execute(sql, {'embedding': embedding_str})
            
            documents = []
            for row in result:
                documents.append({
                    'id': str(row.id),
                    'filename': row.filename,
                    'title': row.title,
                    'summary': row.summary or '',
                    'product_tags': row.product_tags or [],
                    'project_number': row.project_number,
                    'main_city': row.main_city,
                    'main_companies': row.main_companies or [],
                    'file_type': row.file_type,
                    'file_size': row.file_size,
                    'modification_date': row.modification_date,
                    'main_tag': row.main_tag,
                    'vector_distance': float(row.vector_distance) if row.vector_distance else 1.0,
                    'star_count': row.star_count or 0,
                    'contribution_count': row.contribution_count or 0
                })
            
            self.db.commit()
            return documents
            
        except Exception as e:
            logger.error(f"Error finding similar documents: {e}")
            self.db.rollback()
            return []
    
    def _calculate_similarity_scores(
        self,
        documents: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        embedding: List[float]
    ) -> List[Dict[str, Any]]:
        """Calculate multi-factor similarity scores"""
        scored_docs = []
        
        for doc in documents:
            # Vector similarity (inverse of distance, normalized)
            vector_score = 1 / (1 + doc['vector_distance'])
            
            # Product matching score
            product_matches = set(doc['product_tags']) & set(metadata['products'])
            product_score = len(product_matches) / max(len(metadata['products']), 1) if metadata['products'] else 0
            
            # SAU matching score
            sau_match = 0
            if metadata['sau_numbers'] and doc['project_number']:
                if any(sau in doc['project_number'] for sau in metadata['sau_numbers']):
                    sau_match = 1
            
            # Location matching score
            location_match = 0
            if metadata['main_city'] and doc.get('main_city'):
                if metadata['main_city'].lower() in doc['main_city'].lower():
                    location_match = 1
            
            # Company matching score
            company_match = 0
            if metadata['main_company'] and doc.get('main_companies'):
                if any(metadata['main_company'].lower() in comp.lower() 
                    for comp in doc['main_companies']):
                    company_match = 1
            
            if metadata.get('sau_numbers'):
                # Use normal weights when SAU is present
                weights = {
                    'vector': 0.5,
                    'products': 0.3,
                    'sau': 0.1,
                    'city': 0.05,
                    'company': 0.05
                }
            else:
                # Redistribute SAU weight to vector similarity when no SAU
                weights = {
                    'vector': 0.6,  # Added the 0.1 from SAU
                    'products': 0.3,
                    'sau': 0,
                    'city': 0.05,
                    'company': 0.05
                }

            # Then calculate the final score:
            final_score = (
                vector_score * weights['vector'] +
                product_score * weights['products'] +
                sau_match * weights['sau'] +
                location_match * weights['city'] +
                company_match * weights['company']
            )
            
            # Get relative date
            if doc['modification_date']:
                relative_date = self._get_relative_date(doc['modification_date'])
                # Convert datetime to ISO string for JSON serialization
                if isinstance(doc['modification_date'], datetime):
                    modification_date_str = doc['modification_date'].isoformat()
                else:
                    modification_date_str = str(doc['modification_date'])
            else:
                relative_date = 'Date unknown'
                modification_date_str = None
            
            scored_docs.append({
                'id': doc['id'],
                'filename': doc['filename'],
                'title': doc['title'] or doc['filename'],
                'summary': doc['summary'][:200] + '...' if len(doc['summary']) > 200 else doc['summary'],
                'similarity_score': final_score,
                'matching_products': list(product_matches),
                'matching_sau': doc['project_number'] if sau_match else None,
                'matching_city': doc.get('main_city') if location_match else None,
                'matching_company': doc['main_companies'][0] if company_match and doc['main_companies'] else None,
                'file_type': doc['file_type'],
                'file_size': doc['file_size'],
                'modification_date': modification_date_str,  # Use string version
                'relative_date': relative_date,
                'is_starred': doc['star_count'] > 0,
                'contributions_count': doc['contribution_count']
            })
        
        # Sort by final score
        scored_docs.sort(key=lambda x: x['similarity_score'], reverse=True)
        
        return scored_docs
    
    def _get_relative_date(self, date) -> str:
        """Get relative date string"""
        if not date:
            return 'Date unknown'
        
        # Import timezone utilities
        from datetime import timezone
        
        # Get current time as offset-aware (UTC)
        now = datetime.now(timezone.utc)
        
        # Handle string dates
        if isinstance(date, str):
            try:
                date = datetime.fromisoformat(date)
            except:
                return 'Date unknown'
        
        # Make sure date is offset-aware
        if date.tzinfo is None:
            # If date is naive, assume it's UTC
            date = date.replace(tzinfo=timezone.utc)
        
        # Now both are offset-aware and we can subtract
        diff = now - date
        
        if diff.days == 0:
            return 'Today'
        elif diff.days == 1:
            return 'Yesterday'
        elif diff.days < 7:
            return f'{diff.days} days ago'
        elif diff.days < 30:
            weeks = diff.days // 7
            return f'{weeks} week{"s" if weeks > 1 else ""} ago'
        elif diff.days < 365:
            months = diff.days // 30
            return f'{months} month{"s" if months > 1 else ""} ago'
        else:
            years = diff.days // 365
            return f'{years} year{"s" if years > 1 else ""} ago'
    
    async def get_analysis_result(self, task_id: str, user_id: int) -> Optional[Dict[str, Any]]:
        """Get analysis result by task ID"""
        if task_id not in self.analysis_tasks:
            return None
        
        task = self.analysis_tasks[task_id]
        
        if task["user_id"] != user_id:
            return None
        
        if task["status"] == "completed":
            return task["result"]
        
        return {
            "status": task["status"],
            "message": "Analysis in progress"
        }
    
    async def get_user_history(
        self, 
        user_id: int, 
        limit: int = 10,
        offset: int = 0
    ) -> List[Dict[str, Any]]:
        """Get user's analysis history"""
        # In production, store in database
        user_tasks = [
            {
                "task_id": task_id,
                "file_name": task["file_name"],
                "started_at": task["started_at"].isoformat(),
                "status": task["status"],
                "result": task.get("result")
            }
            for task_id, task in self.analysis_tasks.items()
            if task["user_id"] == user_id and task["status"] == "completed"
        ]
        
        # Sort by date
        user_tasks.sort(key=lambda x: x["started_at"], reverse=True)
        
        # Paginate
        return user_tasks[offset:offset + limit]